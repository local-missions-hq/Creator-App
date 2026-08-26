from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "business-plan"
ASSET_DIR = OUT_DIR / "assets"
OUTPUT = OUT_DIR / "Local_Missions_Investor_Business_Plan_2026.docx"
COVER_IMAGE = ASSET_DIR / "local-missions-cover-hero.png"
CREATOR_OVERVIEW = ROOT / "ux-walkthrough" / "creator-overview.png"
BUSINESS_OVERVIEW = ROOT / "ux-walkthrough" / "business-overview.png"


NAVY = "102A43"
LAGOON = "007C83"
TANGERINE = "CF3F1F"
SAND = "FFF7ED"
PALM = "137A50"
GOLD = "D97706"
SLATE = "526273"
BORDER = "E5D8C8"
LIGHT_BLUE = "EAF4F5"
LIGHT_ORANGE = "FDF0EA"
LIGHT_GREEN = "EAF5EF"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"
BLACK = "17212B"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def hex_fill(hex_value: str) -> tuple[int, int, int]:
    return tuple(int(hex_value[i : i + 2], 16) for i in (0, 2, 4))


def draw_market_chart(path: Path) -> None:
    width, height = 1600, 900
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((90, 55), "U.S. creator advertising is becoming a core media channel", font=pil_font(48, True), fill=hex_fill(NAVY))
    d.text((90, 125), "IAB reported / projected spend, $ billions", font=pil_font(28), fill=hex_fill(SLATE))

    years = ["2021", "2024", "2025", "2026E"]
    values = [13.9, 29.5, 37.0, 44.0]
    colors = [SLATE, LAGOON, TANGERINE, NAVY]
    chart_left, chart_top, chart_right, chart_bottom = 140, 220, 1490, 760
    max_val = 50
    for tick in range(0, 51, 10):
        y = chart_bottom - (tick / max_val) * (chart_bottom - chart_top)
        d.line((chart_left, y, chart_right, y), fill=(223, 229, 234), width=2)
        d.text((75, y - 15), f"{tick}", font=pil_font(22), fill=hex_fill(SLATE))
    slot = (chart_right - chart_left) / len(values)
    for idx, (year, value, color) in enumerate(zip(years, values, colors)):
        x0 = chart_left + idx * slot + 95
        x1 = chart_left + (idx + 1) * slot - 95
        y0 = chart_bottom - (value / max_val) * (chart_bottom - chart_top)
        d.rounded_rectangle((x0, y0, x1, chart_bottom), radius=16, fill=hex_fill(color))
        d.text(((x0 + x1) / 2, y0 - 48), f"${value:.1f}B", anchor="mm", font=pil_font(30, True), fill=hex_fill(color))
        d.text(((x0 + x1) / 2, chart_bottom + 38), year, anchor="mm", font=pil_font(27, True), fill=hex_fill(NAVY))
    d.text((90, 835), "Source: IAB 2025 Creator Economy Ad Spend & Strategy Report. 2026 is an IAB expectation, not realized spend.", font=pil_font(21), fill=hex_fill(SLATE))
    img.save(path)


def draw_network_chart(path: Path) -> None:
    width, height = 1600, 1000
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 45), "The defensible asset is local liquidity, not a national creator directory", font=pil_font(45, True), fill=hex_fill(NAVY))
    d.text((80, 110), "Cross-side effects compound inside a city cell; quality controls prevent congestion from destroying value.", font=pil_font(27), fill=hex_fill(SLATE))

    nodes = [
        ((800, 255), "More funded\nlocal missions", TANGERINE),
        ((1260, 500), "More creator\nearnings + retention", LAGOON),
        ((800, 745), "Faster fill +\nbetter mission fit", PALM),
        ((340, 500), "Higher business\nrepeat + referrals", NAVY),
    ]
    for idx, (center, text, color) in enumerate(nodes):
        next_center = nodes[(idx + 1) % len(nodes)][0]
        d.line((center[0], center[1], next_center[0], next_center[1]), fill=hex_fill(GOLD), width=12)
        # arrowhead
        mx = center[0] * 0.35 + next_center[0] * 0.65
        my = center[1] * 0.35 + next_center[1] * 0.65
        d.polygon([(mx, my), (mx - 25, my - 6), (mx - 7, my + 23)], fill=hex_fill(GOLD))
    for center, text, color in nodes:
        box = (center[0] - 210, center[1] - 92, center[0] + 210, center[1] + 92)
        d.rounded_rectangle(box, radius=30, fill=hex_fill(color))
        lines = text.split("\n")
        d.multiline_text(center, "\n".join(lines), anchor="mm", align="center", spacing=8, font=pil_font(30, True), fill="white")

    d.ellipse((585, 380, 1015, 620), fill=hex_fill(SAND), outline=hex_fill(BORDER), width=5)
    d.multiline_text((800, 500), "LOCAL LIQUIDITY\nINDEX", anchor="mm", align="center", spacing=10, font=pil_font(39, True), fill=hex_fill(NAVY))
    d.text((800, 575), "fill × completion × repeat", anchor="mm", font=pil_font(24), fill=hex_fill(SLATE))

    d.rounded_rectangle((115, 855, 1485, 945), radius=20, fill=hex_fill(LIGHT_ORANGE))
    d.text((800, 900), "Expansion rule: do not open the next city until one Orlando cell sustains liquidity, completion, repeat, and contribution-margin gates.", anchor="mm", font=pil_font(25, True), fill=hex_fill(TANGERINE))
    img.save(path)


def draw_unit_economics_chart(path: Path) -> None:
    width, height = 1600, 900
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "Illustrative $1,000 mission: value flows to creators first", font=pil_font(47, True), fill=hex_fill(NAVY))
    d.text((80, 120), "Base case assumes payment costs are passed through at cost; platform value capture is the coordination fee.", font=pil_font(26), fill=hex_fill(SLATE))

    boxes = [
        ("Creator reward pool", "$1,000", PALM, "10 rewards + optional\namplification bonuses"),
        ("Platform net revenue", "$180", LAGOON, "18% coordination,\nverification, rights,\nreporting"),
        ("Variable platform ops", "($52)", GOLD, "support, fraud reserve,\ncloud/media, verification"),
        ("Contribution profit", "$128", TANGERINE, "71% of net revenue;\npilot target ≥55%"),
    ]
    start_x, gap, box_w, box_h = 80, 35, 335, 350
    for idx, (label, value, color, detail) in enumerate(boxes):
        x0 = start_x + idx * (box_w + gap)
        y0 = 240
        d.rounded_rectangle((x0, y0, x0 + box_w, y0 + box_h), radius=26, fill=hex_fill(color))
        d.text((x0 + 28, y0 + 35), label, font=pil_font(25, True), fill="white")
        d.text((x0 + 28, y0 + 112), value, font=pil_font(56, True), fill="white")
        d.multiline_text((x0 + 28, y0 + 220), detail, font=pil_font(21), fill="white", spacing=8)
        if idx < len(boxes) - 1:
            ax = x0 + box_w + 7
            ay = y0 + box_h / 2
            d.polygon([(ax, ay - 22), (ax + 25, ay), (ax, ay + 22)], fill=hex_fill(BORDER))
    d.rounded_rectangle((150, 670, 1450, 815), radius=24, fill=hex_fill(SAND), outline=hex_fill(BORDER), width=4)
    d.text((800, 712), "Business cash invoice: $1,240", anchor="mm", font=pil_font(34, True), fill=hex_fill(NAVY))
    d.text((800, 760), "$1,000 creator rewards + $180 platform fee + ~$60 payment-cost pass-through; any meal/experience is provided in-kind.", anchor="mm", font=pil_font(23), fill=hex_fill(SLATE))
    d.text((80, 855), "Planning model only. Processor fees, refunds, taxes, support burden, and final pricing must be proven in the pilot.", font=pil_font(21), fill=hex_fill(SLATE))
    img.save(path)


def draw_financial_chart(path: Path) -> None:
    width, height = 1600, 900
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "Base case: density first, then measured city replication", font=pil_font(47, True), fill=hex_fill(NAVY))
    d.text((80, 120), "Illustrative management case; $ millions", font=pil_font(27), fill=hex_fill(SLATE))

    years = ["Y1", "Y2", "Y3", "Y4", "Y5"]
    revenue = [0.05, 0.44, 2.11, 8.75, 29.33]
    gmv = [0.30, 1.89, 9.20, 37.50, 129.60]
    chart_left, chart_top, chart_right, chart_bottom = 130, 220, 1480, 760
    max_gmv = 140
    for tick in range(0, 141, 20):
        y = chart_bottom - tick / max_gmv * (chart_bottom - chart_top)
        d.line((chart_left, y, chart_right, y), fill=(226, 231, 235), width=2)
        d.text((65, y - 12), str(tick), font=pil_font(20), fill=hex_fill(SLATE))
    slot = (chart_right - chart_left) / len(years)
    points = []
    for idx, (year, rev, gm) in enumerate(zip(years, revenue, gmv)):
        cx = chart_left + idx * slot + slot / 2
        bar_w = 112
        y0 = chart_bottom - gm / max_gmv * (chart_bottom - chart_top)
        d.rounded_rectangle((cx - bar_w / 2, y0, cx + bar_w / 2, chart_bottom), radius=14, fill=hex_fill(LIGHT_BLUE), outline=hex_fill(LAGOON), width=3)
        d.text((cx, y0 - 26), f"${gm:.1f}", anchor="mm", font=pil_font(22, True), fill=hex_fill(LAGOON))
        ry = chart_bottom - rev / max_gmv * (chart_bottom - chart_top)
        points.append((cx, ry))
        d.text((cx, chart_bottom + 38), year, anchor="mm", font=pil_font(24, True), fill=hex_fill(NAVY))
    d.line(points, fill=hex_fill(TANGERINE), width=10, joint="curve")
    for (cx, ry), rev in zip(points, revenue):
        d.ellipse((cx - 12, ry - 12, cx + 12, ry + 12), fill=hex_fill(TANGERINE))
        d.text((cx + 18, ry - 20), f"${rev:.2f}", font=pil_font(20, True), fill=hex_fill(TANGERINE))
    d.rounded_rectangle((100, 805, 1500, 870), radius=18, fill=hex_fill(SAND))
    d.text((800, 838), "Bars = mission GMV  •  Orange line = net revenue  •  Break-even is modeled in Year 5, not assumed in the pilot.", anchor="mm", font=pil_font(23, True), fill=hex_fill(NAVY))
    img.save(path)


def generate_charts() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    draw_market_chart(ASSET_DIR / "market-growth.png")
    draw_network_chart(ASSET_DIR / "network-effects.png")
    draw_unit_economics_chart(ASSET_DIR / "mission-unit-economics.png")
    draw_financial_chart(ASSET_DIR / "five-year-model.png")


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_table_outer_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=SLATE)


def configure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def make_abstract(abs_id: int, fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "279")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), LAGOON if fmt == "bullet" else NAVY)
        r_pr.extend([fonts, color])
        lvl.extend([start, num_fmt, lvl_text, jc, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)

    def make_num(num_id: int, abs_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abs_id))
        num.append(abstract_num_id)
        numbering.append(num)

    make_abstract(next_abs, "bullet", "•")
    make_num(next_num, next_abs)
    make_abstract(next_abs + 1, "decimal", "%1.")
    make_num(next_num + 1, next_abs + 1)
    return next_num, next_num + 1


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def set_image_alt(inline_shape, description: str, title: str = "") -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    if title:
        doc_pr.set("title", title)


def add_hyperlink(paragraph, text: str, url: str, color: str = LAGOON) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, c, underline])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tokens = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, SLATE, 0, 18),
        "Heading 1": (16, NAVY, 18, 10),
        "Heading 2": (13, LAGOON, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05
        style.paragraph_format.keep_with_next = True
        # Remove inherited Word-template paragraph rules from the title block.
        p_pr = style._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is not None:
            p_pr.remove(p_bdr)

    custom_styles = [
        ("Lead", 12.5, NAVY, True, False, 10),
        ("Kicker", 9.5, TANGERINE, True, False, 5),
        ("Caption", 9, SLATE, False, True, 8),
        ("Source Text", 8.5, SLATE, False, False, 4),
        ("Small", 9, SLATE, False, False, 5),
    ]
    for name, size, color, bold, italic, after in custom_styles:
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.font.italic = italic
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True


def configure_header_footer(section) -> None:
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("LOCAL MISSIONS  •  INVESTOR BUSINESS PLAN")
    set_run_font(run, size=8.5, color=SLATE, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = fp.add_run("CONFIDENTIAL WORKING DRAFT")
    set_run_font(left, size=8.5, color=SLATE, bold=True)
    fp.add_run("\t")
    date = fp.add_run("AUGUST 2026  •  ")
    set_run_font(date, size=8.5, color=SLATE)
    add_page_number(fp)


def add_body(doc: Document, text: str, *, style: str | None = None, bold_lead: str | None = None, italic: bool = False, keep_with_next: bool = False) -> object:
    p = doc.add_paragraph(style=style or "Normal")
    p.paragraph_format.keep_with_next = keep_with_next
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead) :])
        set_run_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_bullets(doc: Document, items: Iterable[str], bullet_num_id: int) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        apply_num(p, bullet_num_id)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str], decimal_num_id: int) -> None:
    numbering = doc.part.numbering_part.element
    base_num = next(
        num for num in numbering.findall(qn("w:num"))
        if int(num.get(qn("w:numId"))) == decimal_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    fresh_num_id = max(
        (int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))),
        default=0,
    ) + 1
    fresh_num = OxmlElement("w:num")
    fresh_num.set(qn("w:numId"), str(fresh_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_num_id)
    fresh_num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    fresh_num.append(level_override)
    numbering.append(fresh_num)

    for item in items:
        p = doc.add_paragraph(style="Normal")
        apply_num(p, fresh_num_id)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(item)


def add_callout(doc: Document, label: str, text: str, fill: str = SAND, accent: str = TANGERINE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_row_cant_split(table.rows[0])
    set_table_geometry(table, [9360])
    remove_table_outer_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=190, bottom=150, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(f"{label.upper()}  ")
    set_run_font(r1, size=10.5, color=accent, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=NAVY, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_dxa: Sequence[int], *, font_size: float = 9.2, header_fill: str = NAVY, first_col_bold: bool = False, alignments: Sequence[str] | None = None) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignments and alignments[idx] == "center" else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, color=WHITE, bold=True)
    for r_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        if r_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for c_idx, (cell, value) in enumerate(zip(cells, row_data)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if alignments and alignments[c_idx] == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK, bold=bool(first_col_bold and c_idx == 0))
    set_table_geometry(table, widths_dxa)
    post = doc.add_paragraph()
    post.paragraph_format.space_before = Pt(0)
    post.paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc: Document, path: Path, caption: str, alt_text: str, width: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt(shape, alt_text, caption.split(".")[0])
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_with_next = False
    cp.add_run(caption)


def add_section(doc: Document, number: str, title: str, lead: str | None = None) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(f"{number}  {title}")
    if lead:
        add_body(doc, lead, style="Lead")


def add_source(doc: Document, number: int, citation: str, url: str) -> None:
    p = doc.add_paragraph(style="Source Text")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(f"[{number}] {citation} ")
    set_run_font(r, size=8.5, color=BLACK)
    add_hyperlink(p, url, url)


def build_document() -> None:
    generate_charts()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    section = doc.sections[0]
    bullet_num_id, decimal_num_id = configure_numbering(doc)

    doc.core_properties.title = "Local Missions Investor Business Plan"
    doc.core_properties.subject = "Market thesis, unit economics, network effects, go-to-market, financial model, and exit strategy"
    doc.core_properties.author = "Local Missions"
    doc.core_properties.keywords = "Local Missions, creator economy, local commerce, marketplace, Orlando, business plan"
    doc.core_properties.comments = "Investor working plan; assumptions as of August 25, 2026"
    doc.core_properties.created = datetime(2026, 8, 25, 16, 55)
    doc.core_properties.modified = datetime(2026, 8, 25, 16, 55)

    # Cover: editorial_cover pattern with a brand-specific hero image.
    p = doc.add_paragraph(style="Kicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run("INVESTOR BUSINESS PLAN  |  AUGUST 2026")
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("LOCAL MISSIONS")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("The operating system for paid, measurable local creator activations")
    meta = doc.add_paragraph(style="Small")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Orlando beachhead  •  U.S. expansion thesis  •  Confidential working draft")
    add_figure(
        doc,
        COVER_IMAGE,
        "Original concept image generated for Local Missions. The scene is illustrative and uses synthetic people and businesses.",
        "Diverse adult local creators capture smartphone content at an independent Central Florida venue while the business owner coordinates on a tablet.",
        width=6.5,
    )
    add_callout(doc, "Investor thesis", "Fund the proof, not the pitch: Local Missions becomes venture-scale only if it proves dense local liquidity, repeat business demand, and software-like contribution margins.", fill=LIGHT_ORANGE, accent=TANGERINE)
    legal = doc.add_paragraph(style="Small")
    legal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legal.add_run("Prepared from current market research and the Local Missions product plan. Forecasts are illustrative, not audited, and not legal, tax, accounting, or investment advice.")
    doc.add_page_break()

    # Contents.
    doc.add_paragraph("CONTENTS", style="Kicker")
    doc.add_paragraph("Business plan at a glance", style="Title")
    contents = [
        ("01", "Executive verdict"),
        ("02", "The problem and the wedge"),
        ("03", "Product and value proposition"),
        ("04", "Where Local Missions fits in the economy"),
        ("05", "Market evidence and sizing"),
        ("06", "Customer and creator strategy"),
        ("07", "Competitive landscape"),
        ("08", "Business model and pricing"),
        ("09", "Unit economics"),
        ("10", "Network effects and defensibility"),
        ("11", "Go-to-market and city expansion"),
        ("12", "Operating model, technology, and trust"),
        ("13", "Five-year financial case and capital plan"),
        ("14", "Risks and kill criteria"),
        ("15", "Exit strategies"),
        ("16", "Milestones and investor conclusion"),
        ("A", "Appendix: assumptions, research, and source notes"),
    ]
    add_table(doc, ["Section", "What it answers"], contents, [1050, 8310], font_size=10, header_fill=NAVY, first_col_bold=True, alignments=["center", "left"])
    add_callout(doc, "Reading lens", "This is a venture case with explicit conditions. The plan distinguishes external facts from management assumptions, and it treats the Orlando pilot as an experiment designed to invalidate weak economics quickly.", fill=LIGHT_BLUE, accent=LAGOON)
    doc.add_page_break()

    # 01 Executive verdict.
    add_section(doc, "01", "Executive verdict", "The concept is investable as a disciplined local-commerce marketplace—not as another generic influencer directory.")
    add_callout(doc, "Verdict", "CONVICTION WITH CONDITIONS. Local Missions attacks a real and growing budget category, but the company must prove repeatable local liquidity and keep manual service from turning the platform into a low-margin agency.", fill=LIGHT_ORANGE, accent=TANGERINE)
    add_body(doc, "Local Missions connects location-based businesses with local adults who complete clearly scoped, paid promotional missions. A mission may require a verified visit, original vertical video and photos, an optional disclosed social post, and a measurable action through a Local Pass, QR code, booking link, or redemption. The product coordinates the brief, participant fit, funding status, on-site verification, deliverables, content rights, review, payout tracking, and results reporting in one workflow.")
    add_body(doc, "The strategic wedge is not influencer fame. It is reliable local execution. Businesses often need authentic content and foot traffic but cannot afford an agency, do not know which creators will actually show up, and struggle to connect social activity to store visits or sales. Everyday creators want legitimate paid opportunities without needing a massive audience. Local Missions converts that mismatch into a repeatable marketplace transaction.")
    doc.add_paragraph("Why this can be large", style="Heading 2")
    add_bullets(doc, [
        "Budget tailwind: IAB expects U.S. creator ad spend to reach $44 billion in 2026 after $37 billion in 2025, while buyers still cite creator selection, measurement, and fragmented operations as major gaps. [1]",
        "Customer breadth: the United States has 36.2 million small businesses; Florida alone has about 3.49 million. Local Missions needs only the location-based, mission-suitable subset—not the whole number—to build a meaningful market. [2]",
        "Strong Orlando test bed: Orlando welcomed a record 76.7 million visitors in 2025, and the metro population approached 3.0 million. The market has unusually dense restaurants, attractions, experiences, hospitality, retail, and family-oriented demand. [4][5]",
        "Marketplace economics: creators can join free, while businesses pay for completed value. That asymmetric pricing is consistent with two-sided-platform theory: subsidize the side whose participation most increases value for the paying side. [7][8]",
        "Strategic scarcity: enterprise creator platforms optimize reach and content at scale; gig apps optimize labor fulfillment; local ad products optimize impressions. Few products combine verified in-person participation, licensed UGC, transparent payments, and local attribution for SMB-sized budgets.",
    ], bullet_num_id)
    doc.add_paragraph("The three things that must be true", style="Heading 2")
    add_numbered(doc, [
        "Businesses repeat. A one-off campaign marketplace is a services treadmill. The target is at least 40% business repeat within 90 days during pilot and 60% annual retention before aggressive scaling.",
        "Creators complete reliably. The platform must fill at least 80% of eligible slots within 72 hours and sustain at least 90% verified completion with low dispute rates.",
        "The work becomes software. Mission-level contribution margin must exceed 55% in pilot and trend toward 65%–75% as review, fraud, support, and media operations become standardized.",
    ], decimal_num_id)
    add_table(doc, ["Investor question", "Required answer before scale"], [
        ["Is demand real?", "20+ paying design partners; 200+ funded missions; no reward subsidy masquerading as revenue."],
        ["Is liquidity real?", "≥80% fill within 72 hours in the target cell; ≥90% verified completion."],
        ["Is retention real?", "≥40% 90-day business repeat; ≥45% creator second-mission rate."],
        ["Are margins real?", "≥55% mission contribution margin including support, verification, fraud reserve, and media costs."],
        ["Is attribution useful?", "≥60% of campaigns capture a verifiable downstream action, not only likes or impressions."],
    ], [2200, 7160], font_size=9.4, header_fill=LAGOON, first_col_bold=True)

    # 02 Problem.
    add_section(doc, "02", "The problem and the wedge", "Local creator marketing is growing, but the last mile between a campaign brief and a measurable in-person outcome is still fragmented.")
    doc.add_paragraph("Business pain", style="Heading 2")
    add_body(doc, "A local operator can buy boosted posts, hire an agency, search a national creator marketplace, or message individuals directly. None of those choices reliably solves the complete operating problem: Who is actually local? Who will arrive during the required window? What exactly will they deliver? Which content rights were granted? Was the material connection disclosed? When should payment advance? Did anyone redeem an offer or book after the activation?")
    add_bullets(doc, [
        "Search cost: finding relevant local people is manual, slow, and often driven by vanity metrics.",
        "Execution risk: no-shows, unclear briefs, inconsistent files, and revision disputes consume owner time.",
        "Trust risk: creators fear nonpayment; businesses fear paying before work is complete.",
        "Measurement gap: impressions do not prove a visit, redemption, booking, or purchase.",
        "Rights and compliance gap: usage terms, sponsorship disclosure, and honest-review rules are frequently handled in scattered messages.",
    ], bullet_num_id)
    doc.add_paragraph("Creator pain", style="Heading 2")
    add_body(doc, "Most people will never be full-time influencers, and that is the point. Local Missions treats reliable local participation, content ability, mission fit, and honest execution as productive capacity. A participant can earn a guaranteed reward and receive an experience without needing a public following. Larger or especially relevant audiences can earn an optional, clearly disclosed distribution bonus; follower count is not the universal gate.")
    doc.add_paragraph("The wedge", style="Heading 2")
    add_callout(doc, "Wedge", "Start with visually rich, location-based Orlando businesses that already buy content or promotions and can measure a visit: independent attractions, family experiences, restaurants/cafes, fitness/wellness, and selected local retail.", fill=LIGHT_GREEN, accent=PALM)
    add_body(doc, "The ideal first customer has an average ticket above roughly $25, meaningful repeat or referral value, unused capacity during specific windows, and a product that is naturally documented with photos or vertical video. The initial product should not attempt every local business category. Dense vertical focus improves mission comparability, support playbooks, creator fit, and attribution.")

    # 03 Product.
    add_section(doc, "03", "Product and value proposition", "The product is a trusted workflow from funded brief to verified visit, licensed content, approved reward, and measurable result.")
    add_table(doc, ["Stage", "Creator experience", "Business value", "Proof created"], [
        ["Fund", "Sees guaranteed reward and terms before applying", "Confirms budget before publication", "Funding confirmation and audit event"],
        ["Match", "Applies based on fit, schedule, distance, and skills", "Selects for mission relevance—not popularity alone", "Application, portfolio, consent, decision"],
        ["Visit", "Checks in during the mission window", "Confirms the right person arrived", "Rotating QR/staff code + coarse time/location proof"],
        ["Deliver", "Uploads the agreed original media", "Receives organized, licensed assets", "File counts, versions, brief snapshot, rights record"],
        ["Review", "Gets one clear revision path and support", "Approves against agreed scope", "Decision, reason, revision history"],
        ["Pay", "Tracks Funded → Pending review → Available → Paid", "Releases payment only through the defined workflow", "Ledger, processor events, payout status"],
        ["Measure", "Optional disclosed posting bonus", "Tracks Local Pass, code, booking, redemption, or purchase", "Attributed action and campaign report"],
    ], [1050, 2600, 2800, 2910], font_size=8.6, header_fill=NAVY, first_col_bold=True)
    add_callout(doc, "Payment-language rule", "Local Missions provides escrow-like confidence through transparent states and processor-backed workflows, but it must not call the product “escrow” unless licensed counsel and the payment partner approve that exact structure.", fill=LIGHT_ORANGE, accent=TANGERINE)
    add_figure(doc, CREATOR_OVERVIEW, "Creator journey concept: discover, apply, check in, deliver, revise if needed, and track payout. Concept art—not implementation evidence.", "Overview of the Local Missions creator journey across a series of iPhone screens.", width=6.25)
    add_figure(doc, BUSINESS_OVERVIEW, "Business journey concept: verify, build, fund, select, review, and measure a mission. Concept art—not implementation evidence.", "Overview of the Local Missions business journey across a responsive dashboard.", width=4.3)
    doc.add_paragraph("Product boundaries that protect the thesis", style="Heading 2")
    add_bullets(doc, [
        "No public social feed, swiping, creator-to-creator chat, or popularity leaderboard in the MVP.",
        "No continuous background location tracking; use mission-window-only and privacy-minimized proof.",
        "No requirement for positive reviews. Required posts must be honest and clearly disclose payment, free meals, products, or experiences.",
        "No stored-value wallet, cryptocurrency, or off-ledger payout system.",
        "No second city until the first Orlando cell reaches liquidity, completion, repeat, and contribution-margin gates.",
    ], bullet_num_id)

    # 04 Economy fit.
    add_section(doc, "04", "Where Local Missions fits in the economy", "Local Missions is economic infrastructure at the intersection of four large systems: local commerce, the creator economy, flexible work, and marketplace payments.")
    add_table(doc, ["Economic layer", "What exists today", "Local Missions' role"], [
        ["Creator economy", "Brand sponsorships, UGC production, affiliate commerce, audience monetization", "Opens paid work to everyday local adults and ties content to verified real-world participation."],
        ["SMB advertising", "Search, social ads, local media, agencies, direct outreach", "Packages a measurable local activation at a price and complexity level an independent operator can use."],
        ["Experience / visitor economy", "Attractions, restaurants, events, wellness, hospitality, retail", "Turns perishable local capacity and experiences into content, visits, referrals, and attributable offers."],
        ["Flexible work", "Gigs, freelance creative services, promotional staffing", "Defines a bounded outcome with transparent scope, reward, rights, and payment status rather than hourly open-ended labor."],
        ["Fintech / platform ops", "Marketplace onboarding, payment processing, payouts, ledgering", "Adds trust and auditable state transitions without holding itself out as a bank or licensed escrow provider."],
    ], [1800, 3600, 3960], font_size=9, header_fill=LAGOON, first_col_bold=True)
    doc.add_paragraph("Value creation", style="Heading 2")
    add_body(doc, "For businesses, the platform reduces search, coordination, and measurement costs. For creators, it converts local time, judgment, and content skill into paid opportunities. For consumers, it increases discovery of local experiences through content that originates with actual visits. For the regional economy, it routes a portion of marketing spend to local individuals and gives smaller venues a performance-oriented alternative to large agencies.")
    doc.add_paragraph("Value capture", style="Heading 2")
    add_body(doc, "The company captures a percentage of mission GMV for coordination, verification, workflow, content-rights records, reporting, and payment operations. Later, it can add subscription revenue for multi-location planning, reusable templates, team permissions, asset libraries, advanced reporting, and agency workspaces. The company should not monetize raw location data or identity documents; privacy erosion would damage both trust and strategic value.")
    add_callout(doc, "Strategic identity", "The winning category is “local activation infrastructure,” not “influencer marketplace.” That framing broadens the buyer, focuses the product on outcomes, and reduces dependence on follower-based social metrics.", fill=LIGHT_BLUE, accent=LAGOON)

    # 05 Market.
    add_section(doc, "05", "Market evidence and sizing", "External evidence validates the budget and customer base; the actual addressable market is built bottom-up from mission-suitable locations and frequency assumptions.")
    add_figure(doc, ASSET_DIR / "market-growth.png", "Figure 1. U.S. creator advertising spend. IAB reported $13.9B in 2021, $29.5B in 2024, projected $37B for 2025, and expected $44B in 2026. [1]", "Bar chart showing U.S. creator ad spend growing from 13.9 billion dollars in 2021 to an expected 44 billion dollars in 2026.")
    doc.add_paragraph("Evidence, with the investor interpretation", style="Heading 2")
    add_bullets(doc, [
        "IAB's 2025 study surveyed more than 450 U.S. ad-spend decision-makers and found creators becoming a must-buy channel, while measurement, creator identification, and standardization remained persistent pain points. Local Missions is aimed at those infrastructure gaps, at SMB scale. [1]",
        "Goldman Sachs Research estimated in 2023 that the global creator economy could grow from about $250 billion to $480 billion by 2027. This is a directional forecast—not evidence that Local Missions owns a $480 billion TAM. [3]",
        "SBA's 2026 FAQ counts 36.2 million U.S. small businesses. The 2025 profile lists 3.49 million in Florida. The relevant opportunity is the narrower subset with a physical venue, promotable experience, and repeat marketing need. [2]",
        "Visit Orlando reported 76.7 million visitors in 2025, including 70.3 million domestic and 6.3 million international visitors. City reporting places the 2025 metro population near 2.96 million. The combination creates unusually rich local demand and creator supply. [4][5]",
        "Academic research finds that informative content, trustworthiness, attractiveness, and perceived similarity support trust in influencer-branded posts, which in turn affects brand awareness and purchase intention. Local Missions operationalizes fit and credibility without assuming that follower count alone produces trust. [6]",
    ], bullet_num_id)
    doc.add_paragraph("Bottom-up market model", style="Heading 2")
    add_table(doc, ["Layer", "Planning assumption", "Mission GMV", "Net platform revenue at 18%"], [
        ["Orlando beachhead", "2,500 mission-suitable venues × 6 missions/year × $1,100", "$16.5M", "$3.0M"],
        ["25-metro SAM", "150,000 venues × 6 missions/year × $1,200", "$1.08B", "$194M"],
        ["U.S. TAM", "1.2M venues × 6 missions/year × $1,250", "$9.0B", "$1.62B"],
        ["Year-5 SOM", "8,000 active businesses × 12 missions/year × $1,350", "$129.6M", "$23.3M + $6.0M SaaS/other"],
    ], [1700, 4400, 1500, 1760], font_size=9, header_fill=NAVY, first_col_bold=True, alignments=["left", "left", "center", "center"])
    add_body(doc, "The venue counts above are management assumptions, not published Census totals. Before a priced round, the team should replace them with a reproducible NAICS-by-metro establishment model and a documented eligibility filter. That upgrade matters because investor-quality TAM is a data model, not a large headline number.", italic=True)
    doc.add_paragraph("Why Orlando is a rational beachhead", style="Heading 2")
    add_bullets(doc, [
        "High density of promotable, visual, capacity-sensitive experiences.",
        "Large resident and visitor pools support both local discovery and travel-related content.",
        "Geography is bounded enough to manage service quality while still containing multiple neighborhood cells.",
        "Founder-led partnerships can be concentrated through Main Streets, chambers, tourism networks, agencies, restaurant groups, and local creator communities.",
        "Seasonality provides a useful stress test for demand forecasting and creator availability.",
    ], bullet_num_id)

    # 06 Customer.
    add_section(doc, "06", "Customer and creator strategy", "Build the marketplace around one paying customer profile and one abundant, respected supply profile.")
    doc.add_paragraph("Primary paying customer: the visual local operator", style="Heading 2")
    add_table(doc, ["Attribute", "Pilot sweet spot"], [
        ["Business type", "Independent or small-chain attraction, experience, restaurant/cafe, fitness/wellness, or selected retail."],
        ["Marketing team", "Owner/operator or 1–5 person marketing team; too small for full-time creator operations."],
        ["Economics", "AOV roughly $25+ or meaningful repeat/referral value; can fund $750–$2,500 campaigns."],
        ["Need", "Fresh vertical content, verified visits, event coverage, new offer, seasonal capacity, or Local Pass redemption."],
        ["Buying trigger", "Opening, launch, menu/service refresh, weekday traffic gap, seasonal event, or paid-media creative fatigue."],
        ["Disqualifier", "Cannot define deliverables, will not fund rewards, requires positive reviews, or operates in a prohibited/safety-sensitive category."],
    ], [1900, 7460], font_size=9.3, header_fill=LAGOON, first_col_bold=True)
    doc.add_paragraph("Primary supply: the reliable local creator", style="Heading 2")
    add_body(doc, "The base mission should be open to adults who meet eligibility and fit requirements, regardless of follower count. Rank for location radius, availability, interests, mission-specific skills, portfolio, reliability, past completion, disclosure behavior, and content quality. A business should see a coarse locality badge—such as “Orlando-area verified” or an approximate service radius—not a home address or exact ZIP code.")
    add_table(doc, ["Creator layer", "Eligibility", "Compensation logic"], [
        ["Everyday UGC participant", "Adult, local fit, available, able to complete the brief", "Guaranteed base reward + included experience; no follower minimum."],
        ["Skilled local creator", "Stronger portfolio or mission-specific production ability", "Higher base reward for scope/quality, not audience size alone."],
        ["Distribution add-on", "Relevant local audience and verified posting capability", "Optional paid bonus for a clearly disclosed post or whitelisted distribution."],
    ], [2100, 3600, 3660], font_size=9, header_fill=NAVY, first_col_bold=True)
    add_callout(doc, "Marketplace ethic", "Do not disguise unpaid promotion as opportunity. If a business requires content or posting, the participant receives clear compensation and sees the material terms before applying.", fill=LIGHT_GREEN, accent=PALM)

    # 07 Competition.
    add_section(doc, "07", "Competitive landscape", "The competition is not one company. It is a stack of partial substitutes that leave the local execution problem unsolved.")
    add_table(doc, ["Category / examples", "Strength", "Structural gap versus Local Missions"], [
        ["Creator marketplaces (Collabstr)", "Large searchable supply, payments, campaign tools, analytics", "Optimized for creator hiring and social reach; local check-in, venue workflow, and offline attribution are not the core product."],
        ["UGC platforms (Insense, Billo)", "Fast licensed content production and optional social distribution", "Product-shipping/content workflow rather than verified local attendance and repeat venue activation."],
        ["Enterprise suites (Sprout/Tagger, Later/Mavely, Influential)", "Scale, enterprise data, campaign management, social commerce", "Pricing and workflow target larger brands/agencies; local SMB activation is not the primary wedge."],
        ["Gig marketplaces (Taskrabbit and staffing tools)", "Local labor fulfillment and trust mechanisms", "Do not combine brand-safe creative briefs, content rights, disclosures, attribution, and creator distribution."],
        ["Local ad/deal products (Meta, Google, Yelp, Groupon)", "Demand generation, discovery, promotions, audience reach", "Sell media or deals, not verified visits plus original creator assets and payout workflow."],
        ["Agencies / direct DMs", "Flexible, relationship-driven, can deliver high-touch strategy", "Expensive or operationally fragmented; hard to standardize, audit, and measure at SMB budgets."],
    ], [2200, 3200, 3960], font_size=8.65, header_fill=NAVY, first_col_bold=True)
    doc.add_paragraph("Pricing benchmark", style="Heading 2")
    add_body(doc, "Current public pricing shows customers already accept a combination of subscription and transaction fees. Collabstr lists a 10% hiring fee on its free and Pro tiers and 5% on Premium; Pro is $249/month billed annually. Insense lists a 20% marketplace fee in trial, 10% in its brand plan, and subscriptions starting around $500/month. These are not perfect comparables, but they support an 18% pilot take rate for a transaction that adds local verification and operations—provided the workflow truly delivers more value. [12][13]")
    doc.add_paragraph("Positioning map", style="Heading 2")
    add_table(doc, ["Dimension", "Generic creator marketplace", "Local Missions"], [
        ["Unit of value", "Creator profile or post", "Completed, verified local mission"],
        ["Primary buyer", "Brand/agency", "Location-based SMB, multi-location operator, or local agency"],
        ["Selection", "Audience, niche, portfolio, rate", "Local fit, availability, reliability, skills, portfolio; audience is optional"],
        ["Proof", "Content delivered / post analytics", "Funding, check-in, assets, rights, decisions, payout state, redemption/booking"],
        ["Moat", "Creator directory and workflow", "Dense local liquidity + completion data + venue/creator operating history"],
    ], [1700, 3500, 4160], font_size=9.1, header_fill=LAGOON, first_col_bold=True)

    # 08 Business model.
    add_section(doc, "08", "Business model and pricing", "Charge the business for completed coordination value; keep creator access free; add subscriptions only after repeat usage is proven.")
    add_table(doc, ["Revenue stream", "Pilot / launch", "Later-stage expansion", "Investor logic"], [
        ["Mission coordination fee", "18% of creator reward GMV", "15%–20% by volume/service level", "Core transaction revenue; benchmarked within current marketplace ranges."],
        ["Payment-cost pass-through", "Estimated and reconciled at cost", "Negotiated processor economics", "Prevents processor expense from hiding weak platform margin."],
        ["Business Pro", "Not required in pilot", "$99–$299/month for templates, teams, asset library, analytics, and multi-location tools", "Adds recurring revenue only when workflow recurrence exists."],
        ["Agency / enterprise workspace", "Design-partner pilots", "Custom annual contracts + lower transaction rate", "Aggregates demand and improves city liquidity."],
        ["Distribution / paid reach add-on", "Explicit creator bonus", "Workflow and management fee", "Separates base UGC participation from audience amplification."],
        ["Attribution / Local Pass", "Included baseline reporting", "Advanced measurement or performance fee where legally and operationally sound", "Links content to local commerce without selling personal data."],
    ], [1800, 2600, 3300, 1660], font_size=8.5, header_fill=NAVY, first_col_bold=True)
    doc.add_paragraph("Pricing principles", style="Heading 2")
    add_bullets(doc, [
        "Show the business the full invoice: creator rewards, platform fee, payment-cost pass-through, taxes, and any optional distribution bonus.",
        "Show the creator the guaranteed reward and optional bonus separately before application; do not shrink the displayed reward with surprise platform deductions.",
        "Price content scope and usage rights explicitly. A 30-day organic usage license is not the same product as perpetual paid-media rights.",
        "Keep the first mission easy to buy, but never subsidize the reward pool to manufacture GMV. Discounts should apply to the platform fee, not creator compensation.",
        "Move high-frequency operators to subscription + lower take-rate plans only when the economics improve retention and total contribution profit.",
    ], bullet_num_id)

    # 09 Unit economics.
    add_section(doc, "09", "Unit economics", "A venture marketplace is not saved by GMV. It is saved by net revenue, repeat behavior, contribution margin, and declining acquisition/support costs as density improves.")
    add_figure(doc, ASSET_DIR / "mission-unit-economics.png", "Figure 2. Illustrative mission unit economics. Payment costs are modeled as a transparent pass-through; creator rewards are GMV, not platform revenue.", "Diagram of a one-thousand-dollar creator reward pool, a one-hundred-eighty-dollar platform fee, fifty-two dollars of variable operations, and one-hundred-twenty-eight dollars of contribution profit.")
    add_table(doc, ["Base mission invoice", "Amount", "Treatment"], [
        ["Creator reward pool", "$1,000", "Marketplace GMV / creator payable"],
        ["Platform coordination fee (18%)", "$180", "Net platform revenue"],
        ["Estimated payment-cost pass-through", "$60", "Offset to actual processing/Connect/payout costs"],
        ["Business cash invoice", "$1,240", "Excludes any in-kind meal or experience"],
        ["Variable platform operations", "$52", "Support $30; verification/fraud reserve $10; cloud/media/notifications $12"],
        ["Contribution profit", "$128", "71% of net platform revenue; 10.3% of business cash invoice"],
    ], [4200, 1500, 3660], font_size=9.2, header_fill=LAGOON, first_col_bold=True, alignments=["left", "center", "left"])
    add_body(doc, "Stripe currently lists $2 per monthly active account plus 0.25% + $0.25 per payout when the platform handles pricing, with other payment and risk costs depending on the integration. A 10-creator mission can therefore carry meaningful per-account and payout expense before card processing, refunds, disputes, or instant payouts. The pilot must reconcile the pass-through to actual processor statements rather than rely on a blended estimate. [10]")
    doc.add_paragraph("Business acquisition and lifetime value", style="Heading 2")
    add_table(doc, ["Metric", "Conservative pilot case", "Density target case"], [
        ["Business CAC", "$450", "$300"],
        ["Missions per active business / year", "6", "8"],
        ["Average mission GMV", "$1,000", "$1,000"],
        ["Net platform revenue / mission", "$180", "$180"],
        ["Contribution margin", "60%", "68%"],
        ["Annual business retention", "60%", "70%"],
        ["Annual contribution / business", "$648", "$979"],
        ["Contribution LTV (annual contribution ÷ churn)", "$1,620", "$3,263"],
        ["LTV / CAC", "3.6×", "10.9×"],
        ["CAC payback", "4.2 missions / ~8 months", "2.5 missions / ~4 months"],
    ], [3550, 2905, 2905], font_size=9, header_fill=NAVY, first_col_bold=True, alignments=["left", "center", "center"])
    add_callout(doc, "Scale gate", "Do not buy growth until the trailing cohort shows ≥3× contribution LTV/CAC, <12-month payback, and repeat demand that is not founder-dependent.", fill=LIGHT_ORANGE, accent=TANGERINE)
    doc.add_paragraph("What moves the economics", style="Heading 2")
    add_bullets(doc, [
        "Mission frequency is more valuable than a higher first-order take rate. Six repeat missions produce better LTV than one high-fee campaign.",
        "Creator retention lowers supply acquisition, onboarding, and failure costs. The second-mission rate is an economic KPI, not a vanity metric.",
        "Standard briefs and rights templates reduce support minutes and disputes.",
        "Local density lowers both business CAC and time-to-fill; national breadth without local depth does the opposite.",
        "Media volume, chargebacks, manual review, and insurance can destroy margin if treated as overhead instead of mission-level variable cost.",
    ], bullet_num_id)

    # 10 Network effects.
    add_section(doc, "10", "Network effects and defensibility", "The platform can develop real network effects, but they are local, operational, and earned—not automatic.")
    add_figure(doc, ASSET_DIR / "network-effects.png", "Figure 3. The Local Missions cross-side flywheel. More funded missions improve creator earnings and retention; reliable supply improves fill and fit; better execution increases business repeat.", "Flywheel showing funded missions, creator retention, faster fill and fit, and business repeat around a local liquidity index.")
    add_body(doc, "Rochet and Tirole's foundational model describes platforms that must get both sides on board, and Evans and Schmalensee show why one side may be priced below marginal cost when its participation raises the value of the other side. Local Missions should keep creator access free because creator participation increases selection, speed, and reliability for the business that pays. [7][8]")
    doc.add_paragraph("Four defensibility layers", style="Heading 2")
    add_numbered(doc, [
        "Cross-side local network effects. More quality missions attract reliable local creators; better local supply improves fill and outcomes; better outcomes increase business repeat and referrals.",
        "Operational data effects. Completed missions produce reliability, check-in, delivery, revision, rights, dispute, and conversion signals. Those signals improve matching and risk controls without needing invasive personal data.",
        "Workflow switching costs. Businesses accumulate reusable briefs, creator histories, approved assets, rights windows, Local Pass performance, and multi-location templates. Creators accumulate trusted completion and payout history.",
        "Distribution and partnership advantage. Main Street groups, local agencies, venue networks, and creator communities can become proprietary acquisition channels if performance is better than direct coordination.",
    ], decimal_num_id)
    doc.add_paragraph("Why network effects can fail", style="Heading 2")
    add_table(doc, ["Failure mode", "Economic mechanism", "Countermeasure"], [
        ["Thin local supply", "National creator count does not help a mission that needs ten people near one venue Wednesday afternoon", "Launch neighborhood/vertical cells; track eligible active supply, not registered accounts."],
        ["Creator congestion", "Too many creators chasing too few missions reduces earnings and retention", "Cap acquisition by cell; waitlists; transparent eligibility; reward reliable repeat participation."],
        ["Business multi-homing", "Businesses can post to several platforms or DM creators", "Win on speed, proof, rights, payout trust, attribution, and saved workflows—not exclusivity tricks."],
        ["Off-platform leakage", "Repeat pairs may bypass the fee", "Make repeat workflow, insurance/support, audit history, payment tracking, and attribution worth more than the take rate."],
        ["Manual-service trap", "More GMV requires proportional coordinators", "Standardize briefs, automation, exception queues, and partner enablement; measure support minutes per mission."],
    ], [1850, 3550, 3960], font_size=8.7, header_fill=LAGOON, first_col_bold=True)
    add_body(doc, "Harvard research on platform interconnectivity distinguishes global marketplaces from local network clusters such as ride-sharing, group buying, local services, and reservations. Local Missions belongs to the local-cluster class: the moat is density within each market, with playbooks and technology creating some intercity leverage. [9]")

    # 11 GTM.
    add_section(doc, "11", "Go-to-market and city expansion", "Sequence demand, supply, and geography so the company creates liquidity instead of collecting dormant accounts.")
    doc.add_paragraph("Phase 1: Orlando design-partner cell", style="Heading 2")
    add_numbered(doc, [
        "Recruit 20 anchor businesses in one or two adjacent verticals and two compact geographic cells. Secure signed pilot terms, not letters of vague interest.",
        "Recruit 300–500 eligible adult creators from local communities, referrals, colleges, hospitality/creative networks, and existing creator groups—without overloading the cell.",
        "Run 50 white-glove missions to discover failure modes and price support. Then run 150 additional missions with progressively standardized workflows.",
        "Publish a cohort scorecard: fill time, completion, submission quality, support minutes, disputes, creator repeat, business repeat, attributed actions, and contribution margin.",
        "Convert the strongest design partners to standard pricing and referrals. Do not report subsidized pilots as proof of willingness to pay.",
    ], decimal_num_id)
    doc.add_paragraph("Demand channels", style="Heading 2")
    add_table(doc, ["Channel", "Offer", "Reason it can work"], [
        ["Founder-led outbound", "Campaign blueprint + transparent $1,000 example", "Fast feedback and control over the first customer profile."],
        ["Main Street / chamber / tourism groups", "Member workshop + vetted pilot cohort", "Aggregated local trust and lower CAC."],
        ["Local agencies", "White-label/partner workspace and creator operations", "Concentrated repeat demand without replacing strategic agency work."],
        ["Venue and franchise groups", "Multi-location templates and benchmark reporting", "Replicable missions and higher frequency."],
        ["Creator referrals", "Priority or referral bonus after verified completion", "High-trust supply growth tied to actual marketplace quality."],
    ], [1900, 3500, 3960], font_size=9, header_fill=NAVY, first_col_bold=True)
    doc.add_paragraph("City expansion gate", style="Heading 2")
    add_table(doc, ["Metric", "Minimum for 8 consecutive weeks"], [
        ["Eligible-slot fill within 72 hours", "≥80%"],
        ["Verified completion", "≥90%"],
        ["On-time acceptable submission", "≥85%"],
        ["90-day business repeat", "≥40%"],
        ["Creator second-mission rate", "≥45%"],
        ["Dispute rate", "<3%"],
        ["Mission contribution margin", "≥55%"],
        ["Support time", "<30 minutes per completed mission on median"],
    ], [5200, 4160], font_size=9.3, header_fill=LAGOON, first_col_bold=True, alignments=["left", "center"])
    add_callout(doc, "Expansion rule", "Clone the playbook, not the chaos. The second market opens only after Orlando works without heroic founder intervention.", fill=LIGHT_GREEN, accent=PALM)

    # 12 Operations/tech/trust.
    add_section(doc, "12", "Operating model, technology, and trust", "Trust is not a legal page. It is the product architecture, state machine, evidence trail, and exception handling.")
    doc.add_paragraph("Operating model", style="Heading 2")
    add_table(doc, ["Capability", "MVP approach", "Scale approach"], [
        ["Business approval", "Manual verification and category review", "Risk scoring, document/provider checks, exception queue"],
        ["Mission approval", "Template + admin review", "Policy engine, prohibited-term detection, vertical templates"],
        ["Matching", "Eligibility filters + business selection", "Reliability/fit ranking with fairness monitoring"],
        ["Check-in", "Rotating venue QR + staff fallback; mission-window-only location", "Risk-based signals and venue device management"],
        ["Media", "Direct uploads, resumable workflow, virus/type/count checks", "Automated QC, rights metadata, lifecycle storage"],
        ["Review / disputes", "One defined revision and support escalation", "SLA queues, evidence bundles, policy-based resolutions"],
        ["Payments", "Stripe Connect test mode; ledger + webhooks", "Reconciliation, reserves, controlled live rollout, finance ops"],
    ], [1900, 3650, 3810], font_size=8.7, header_fill=NAVY, first_col_bold=True)
    doc.add_paragraph("Technology thesis", style="Heading 2")
    add_body(doc, "The existing plan proposes a native-feeling React Native/Expo iPhone app, a Next.js business/admin dashboard, a NestJS API, PostgreSQL, Azure Container Apps and Blob Storage, Service Bus, managed identity/Key Vault, Terraform, OpenTelemetry, and Stripe Connect. That stack is credible, but architecture is not the investment thesis. The minimum valuable product is the full test-mode workflow from funded mission through verified check-in, submission, approval, and reconciled payout—not a polished mission feed.")
    doc.add_paragraph("Trust and regulatory design", style="Heading 2")
    add_bullets(doc, [
        "FTC disclosures: payment, free meals, products, discounts, or experiences are material connections. Disclosures must be clear, conspicuous, and placed with the endorsement; the business must not require a positive review. [11]",
        "Content rights: record permitted channels, duration, territory, edits, paid-media usage, exclusivity, and raw-file ownership before application.",
        "Worker classification: define missions as outcome-based independent work only after state-specific employment counsel reviews control, scheduling, economic dependence, and classification risk.",
        "Payments: use processor-hosted onboarding, KYC, ledgering, reconciliation, reserves, refunds, and disputes. Do not expose identity, tax, or bank data to businesses.",
        "Privacy: collect precise location only when the check-in purpose is immediate; store the minimum evidence; expose only coarse locality to businesses; never sell raw location history.",
        "Safety: prohibit sensitive venues/categories, provide emergency/support paths, verify business locations, and carry insurance appropriate to the operating model.",
    ], bullet_num_id)
    add_callout(doc, "Data moat rule", "Use performance data to improve fit and trust. Do not turn surveillance into the moat. A buyer will value clean consent, rights, and measurement data more than a privacy liability.", fill=LIGHT_BLUE, accent=LAGOON)

    # 13 Financials.
    add_section(doc, "13", "Five-year financial case and capital plan", "The base case assumes a slow Orlando proof year, one replication year, and disciplined multi-city expansion. It is designed to expose the operating leverage required—not to promise the outcome.")
    add_figure(doc, ASSET_DIR / "five-year-model.png", "Figure 4. Illustrative base-case GMV and net revenue. The plan models break-even in Year 5 after four years of investment.", "Chart showing mission GMV rising from 0.3 million dollars in Year 1 to 129.6 million dollars in Year 5, with net revenue rising to 29.33 million dollars.")
    add_table(doc, ["Base case", "Y1", "Y2", "Y3", "Y4", "Y5"], [
        ["Ending active businesses", "75", "300", "1,000", "3,000", "8,000"],
        ["Completed missions", "300", "1,800", "8,000", "30,000", "96,000"],
        ["Average mission GMV", "$1,000", "$1,050", "$1,150", "$1,250", "$1,350"],
        ["Mission GMV", "$0.30M", "$1.89M", "$9.20M", "$37.50M", "$129.60M"],
        ["Transaction net revenue (18%)", "$0.05M", "$0.34M", "$1.66M", "$6.75M", "$23.33M"],
        ["SaaS / other revenue", "$0.00M", "$0.10M", "$0.45M", "$2.00M", "$6.00M"],
        ["Total net revenue", "$0.05M", "$0.44M", "$2.11M", "$8.75M", "$29.33M"],
        ["Gross margin", "55%", "60%", "65%", "70%", "74%"],
        ["Operating expense", "$1.20M", "$2.00M", "$3.80M", "$7.50M", "$15.50M"],
        ["Illustrative EBITDA", "($1.17M)", "($1.74M)", "($2.43M)", "($1.38M)", "$6.20M"],
    ], [2400, 1392, 1392, 1392, 1392, 1392], font_size=8.3, header_fill=NAVY, first_col_bold=True, alignments=["left", "center", "center", "center", "center", "center"])
    doc.add_paragraph("Capital strategy", style="Heading 2")
    add_table(doc, ["Round", "Milestone basis", "Illustrative use"], [
        ["Pre-seed: $1.0M–$1.5M", "Ship the end-to-end test-mode workflow; complete 200+ funded pilot missions; prove ≥55% mission contribution margin", "Product/engineering, trust & safety, pilot operations, counsel/insurance, Orlando partnerships"],
        ["Seed: $3M–$5M", "40%+ 90-day business repeat, 80%+ fill within 72h, <12-month CAC payback, repeatable Orlando acquisition", "Automate operations, deepen Orlando, open 2–4 validated cities, agency/multi-location product"],
        ["Series A or profitable path", "$5M+ ARR, 65%+ gross margin, city playbook works in multiple markets", "Scale distribution, payments economics, enterprise integrations, measurement, selective M&A"],
    ], [1800, 4000, 3560], font_size=8.8, header_fill=LAGOON, first_col_bold=True)
    add_body(doc, "The modeled losses imply roughly $6.7 million of cumulative negative EBITDA before Year 5, excluding working-capital timing, capitalized development, taxes, financing costs, and downside cushion. A staged pre-seed plus seed plan is therefore more credible than a single undercapitalized round. If the pilot misses liquidity or margin gates, the company should narrow the wedge or stop before raising expansion capital.", italic=True)
    doc.add_paragraph("Scenario sensitivity", style="Heading 2")
    add_table(doc, ["Scenario", "Year-5 condition", "Implication"], [
        ["Bear", "3,000 active businesses, six missions/year, support-heavy gross margin below 55%", "Likely $6M–$10M revenue ceiling; reposition as vertical SaaS/managed service or stop geographic expansion."],
        ["Base", "8,000 active businesses, 12 missions/year, 74% gross margin, $29.3M revenue", "Credible strategic acquisition or continued growth case; Year-5 operating profitability."],
        ["Upside", "12,000 active businesses, 15 missions/year, agency/multi-location mix, >75% gross margin", "Potential $50M+ revenue platform with national strategic value; requires exceptional repeat and automation."],
    ], [1400, 4200, 3760], font_size=9, header_fill=NAVY, first_col_bold=True)

    # 14 Risks.
    add_section(doc, "14", "Risks and kill criteria", "A serious plan says what can kill the company and how quickly management will know.")
    add_table(doc, ["Risk", "Early warning", "Mitigation / kill criterion"], [
        ["Chicken-and-egg liquidity", "Fill <60% or >7 days after 100 missions", "Narrow geography/vertical; if two focused cells still fail, stop marketplace expansion."],
        ["Low business repeat", "90-day repeat <25%", "Interview cohorts, tighten ICP, improve measurable offers; kill broad SMB thesis if repeat remains weak."],
        ["Agency economics", ">60 support minutes per mission; gross margin <40%", "Standardize, automate, price service tiers; reject custom work that cannot become product."],
        ["Creator quality / no-shows", "Completion <80%; disputes >5%", "Reliability tiers, deposits only if fair/legal, waitlists, business proof, support; pause risky categories."],
        ["Fraud / payment losses", "Chargebacks and losses exceed reserve", "Risk scoring, delayed availability, limits, reconciliation, processor controls, insurance."],
        ["Platform dependency", "Social API or policy changes break reporting", "Make base value the verified visit and licensed files; treat external posting analytics as an add-on."],
        ["Legal classification / licensing", "Counsel finds worker or money-transmission exposure incompatible with model", "Change control/payment architecture before live launch; never solve legal risk with marketing language."],
        ["Privacy / safety incident", "Unauthorized data access or unsafe mission", "Data minimization, access controls, incident response, prohibited categories, venue verification, insurance."],
        ["AI commoditizes content", "Businesses value synthetic assets over local creators", "Emphasize real visits, local proof, human distribution, rights, and measurable commerce—not generic content volume."],
    ], [1900, 2800, 4660], font_size=8.4, header_fill=NAVY, first_col_bold=True)
    add_callout(doc, "Board discipline", "The company should maintain a written kill-criteria dashboard. Missing one quarter is a problem to diagnose; missing the same structural gate across two focused iterations is evidence, not bad luck.", fill=LIGHT_ORANGE, accent=TANGERINE)

    # 15 Exit strategies.
    add_section(doc, "15", "Exit strategies", "Build a company that can remain independent; preserve multiple strategic exits by owning local liquidity, workflow, and clean attribution data.")
    doc.add_paragraph("The most credible exits", style="Heading 2")
    add_table(doc, ["Path", "Likely buyer logic", "Milestones that create leverage"], [
        ["Creator / social marketing platform", "Adds offline/local activation, SMB access, verified visits, and creator payment workflow", "$15M+ net revenue, 20+ dense markets, strong business repeat, measurable lift, clean rights data"],
        ["SMB operating system / CRM / marketing platform", "Adds a transaction-driven local growth channel to an existing merchant base", "Thousands of retained businesses, multi-location adoption, CRM/booking integrations, <12-month CAC payback"],
        ["Payments / commerce platform", "Adds mission GMV, connected accounts, local attribution, and merchant growth services", "$250M+ annual GMV, low losses, mature reconciliation, clear legal structure"],
        ["Agency / holding company", "Adds proprietary creator supply, local execution, and performance data", "High campaign quality, agency channel revenue, city coverage, data/identity interoperability"],
        ["Private-equity recap / profitable independence", "Durable cash flow, recurring software mix, efficient city expansion", "$20M+ ARR, 15%–25% EBITDA, low concentration, documented operational playbook"],
        ["IPO (not the base case)", "Category-leading local activation infrastructure", "$100M+ net revenue, strong growth, multiple countries/verticals, durable margins and governance"],
    ], [1900, 3400, 4060], font_size=8.6, header_fill=LAGOON, first_col_bold=True)
    doc.add_paragraph("Market signals from recent transactions", style="Heading 2")
    add_bullets(doc, [
        "Sprout Social acquired influencer-marketing and social-intelligence platform Tagger Media for $140 million in 2023, explicitly to expand workflow, reporting, and intelligence. [14]",
        "Later acquired everyday-influencer/social-commerce platform Mavely in a transaction valued at approximately $250 million in January 2025, emphasizing full-funnel creator marketing and measurable ROI. [15]",
        "Publicis Groupe acquired Influential in 2024 to combine creator activation with identity, data, media planning, and measurement. The disclosed strategic logic—not an implied Local Missions valuation—is the relevant signal. [16]",
        "Intuit's approximately $12 billion Mailchimp acquisition illustrates the long-run strategic value of combining SMB operating data with customer acquisition and marketing workflows. It is a strategic adjacency, not a direct comparable. [17]",
    ], bullet_num_id)
    add_callout(doc, "Exit principle", "Do not optimize the company for a logo list of potential acquirers. Optimize for strategic scarcity: dense local supply, retained merchant demand, trusted transaction history, measurable commerce, and a clean legal/data posture.", fill=LIGHT_GREEN, accent=PALM)

    # 16 Milestones.
    add_section(doc, "16", "Milestones and investor conclusion", "The investment case should be re-underwritten at each evidence gate.")
    add_table(doc, ["Window", "Build / sell", "Proof required"], [
        ["0–3 months", "Finalize M0 legal/product decisions; clickable prototype; 20 business discovery partners; recruit founding creator cohort", "Signed pilot criteria, pricing interviews, prohibited-category policy, mission templates, data model"],
        ["4–9 months", "End-to-end dev/test workflow; 50 white-glove missions", "Funding-to-payout audit trail, ≥75% fill, ≥85% completion, processor reconciliation, support-cost baseline"],
        ["10–15 months", "200+ cumulative paid missions; standard pricing; Local Pass attribution", "≥80% fill, ≥90% completion, ≥40% business repeat, ≥45% creator repeat, ≥55% contribution margin"],
        ["16–24 months", "300 active Orlando businesses; agency/multi-location beta; second-city readiness", "<12-month CAC payback, repeatable acquisition, 8-week expansion gate, no founder heroics"],
        ["24–36 months", "2–4 cities; 1,000 active businesses; automation and subscription product", "$2M+ annualized net revenue, 65%+ gross margin, low concentration, stable trust metrics"],
    ], [1500, 3860, 4000], font_size=8.8, header_fill=NAVY, first_col_bold=True)
    doc.add_paragraph("The investor case in one paragraph", style="Heading 2")
    add_body(doc, "Local Missions can become the operating layer that lets a neighborhood business buy authentic local participation as easily as it buys digital ads—while giving everyday adults legitimate paid creator work and giving both sides transparent rules, evidence, and payment states. The market is large enough, the Orlando wedge is unusually strong, and current platform economics support the model. But the company is only venture-scale if repeat demand, local density, and automation compound together. The right investment is a staged bet on proof: build the complete transaction, price it honestly, measure every failure, and expand only when one city cell works.", style="Lead")
    add_callout(doc, "Recommendation", "Proceed to a tightly scoped Orlando pilot and pre-seed raise only after founder/team credibility, legal/payment architecture, and 20 paying design-partner commitments are documented. Make the first board dashboard the liquidity-and-unit-economics scorecard—not downloads or registered users.", fill=LIGHT_ORANGE, accent=TANGERINE)

    # Appendix. It naturally starts on the next page after the conclusion callout.
    add_section(doc, "A", "Appendix: assumptions, research, and source notes", "All external figures are cited. All company forecasts and market layers are management assumptions unless explicitly attributed.")
    doc.add_paragraph("Key management assumptions", style="Heading 2")
    add_table(doc, ["Assumption", "Base case", "Validation method"], [
        ["Platform fee", "18% of creator reward GMV", "Price 20 design partners; test 15% / 18% / 20% conversion and repeat"],
        ["Average mission GMV", "$1,000 in Y1 rising to $1,350 in Y5", "Observe funded cohorts by vertical and creator count"],
        ["Mission frequency", "4 per business in Y1; 12 in Y5", "Measure cohort repeat, seasonality, and subscription adoption"],
        ["Payment pass-through", "~$60 per $1,000 mission", "Reconcile actual processor, Connect, payout, refund, and dispute costs"],
        ["Contribution margin", "55% pilot; 74% blended Y5 gross margin", "Time-track support; allocate fraud, cloud/media, insurance, and verification per mission"],
        ["Business CAC", "$450 pilot; $300 density target", "Fully loaded channel cohort including sales time and partner fees"],
        ["Annual business retention", "60% pilot model; 70% density target", "Logo and GMV retention by acquisition cohort"],
        ["U.S. mission-suitable venues", "1.2 million", "Replace with NAICS-by-metro establishment data and documented eligibility filter"],
    ], [2700, 2850, 3810], font_size=8.8, header_fill=LAGOON, first_col_bold=True)
    doc.add_paragraph("Source notes", style="Heading 2")
    sources = [
        (1, "Interactive Advertising Bureau, 2025 Creator Economy Ad Spend & Strategy Report, published November 20, 2025.", "https://www.iab.com/insights/2025-creator-economy-ad-spend-strategy-report/"),
        (2, "U.S. Small Business Administration Office of Advocacy, Frequently Asked Questions About Small Business (February 2026) and 2025 Small Business Profiles release.", "https://advocacy.sba.gov/wp-content/uploads/2026/02/FINAL_FAQsAboutSmallBusiness_2026_012826.pdf"),
        (3, "Goldman Sachs Research, The creator economy could approach half-a-trillion dollars by 2027, April 19, 2023.", "https://www.goldmansachs.com/insights/articles/the-creator-economy-could-approach-half-a-trillion-dollars-by-2027"),
        (4, "Visit Orlando, Orlando Welcomed Record 76.7 Million Visitors in 2025, May 7, 2026.", "https://www.visitorlando.org/media/press-releases/post/orlando-welcomed-record-767-million-visitors-in-2025-remaining-most-visited-destination-in-the-us/"),
        (5, "City of Orlando, 2025 Popular Annual Financial Report, selected demographic and economic statistics.", "https://www.orlando.gov/Our-Government/Records-and-Documents/Financial/Popular-Financial-Reports/2025-Popular-Annual-Financial-Report"),
        (6, "Lou, C. & Yuan, S. (2019), Influencer Marketing: How Message Value and Credibility Affect Consumer Trust of Branded Content on Social Media, Journal of Interactive Advertising, 19(1), 58–73.", "https://doi.org/10.1080/15252019.2018.1533501"),
        (7, "Rochet, J.-C. & Tirole, J. (2003), Platform Competition in Two-Sided Markets, Journal of the European Economic Association, 1(4), 990–1029.", "https://doi.org/10.1162/154247603322493212"),
        (8, "Evans, D. S. & Schmalensee, R. (2005), The Industrial Organization of Markets with Two-Sided Platforms, NBER Working Paper 11603.", "https://www.nber.org/papers/w11603"),
        (9, "Harvard Business School research, Network Interconnectivity and Entry into Platform Markets, discussion of local platform clusters.", "https://www.hbs.edu/ris/Publication%20Files/Network%20Interconnectivity%20and%20Entry%20into%20Platform%20Markets%20FINAL_78783cb2-0aee-4c1b-8b0d-4fc74e62610f.pdf"),
        (10, "Stripe Connect pricing, accessed August 25, 2026.", "https://stripe.com/connect/pricing"),
        (11, "Federal Trade Commission, Disclosures 101 for Social Media Influencers and Endorsement Guides resources, accessed August 25, 2026.", "https://www.ftc.gov/business-guidance/resources/disclosures-101-social-media-influencers"),
        (12, "Collabstr pricing, accessed August 25, 2026.", "https://collabstr.com/pricing"),
        (13, "Insense pricing, accessed August 25, 2026.", "https://insense.pro/pricing"),
        (14, "Sprout Social, Sprout Social Acquires Tagger Media for $140 million, August 3, 2023.", "https://investors.sproutsocial.com/news/news-details/2023/Sprout-Social-Acquires-Tagger-Media/default.aspx"),
        (15, "Nu Skin Enterprises SEC exhibit, strategic transaction of Mavely for approximately $250 million, January 3, 2025.", "https://www.sec.gov/Archives/edgar/data/1021561/000114036125000113/ef20041017_ex99-1.htm"),
        (16, "Publicis Groupe, Publicis Groupe to Acquire Influential, July 25, 2024.", "https://www.publicisgroupe.com/sites/default/files/press-releases/2024-07/publicis-groupe-to-acquire-influential-creating-world-s-leading-influencer-marketing-solution.pdf"),
        (17, "Intuit SEC filing, completed acquisition of Mailchimp for approximately $12 billion, November 1, 2021.", "https://investors.intuit.com/sec-filings/all-sec-filings/content/0001193125-21-342279/d241331d424b7.htm"),
    ]
    for number, citation, url in sources:
        add_source(doc, number, citation, url)

    doc.add_paragraph("Internal project basis", style="Heading 2")
    add_body(doc, "This business plan also incorporates the Local Missions build contract at plans.md and the synthetic concept walkthrough at ux-walkthrough/README.md in the Creator App repository. Those files define the intended workflow and product boundaries; they are not evidence of implemented production software or live market traction.", style="Source Text")

    configure_sections(doc)
    for sec in doc.sections:
        configure_header_footer(sec)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
